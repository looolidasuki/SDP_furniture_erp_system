"""Section 8.15 Finance Department — UI Design (English), ITP report style: User can..."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

OUTPUT = r"C:\Users\user\source\repos\4915M_claude\FurnitureERP\docs\Section8_15_Finance_Department_UI_Design.docx"


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.1)
    s.right_margin = Inches(1.1)

    title = doc.add_heading("Section 8.15 — Finance Department", level=0)
    for run in title.runs:
        run.font.name = "Calibri"

    para(
        doc,
        "ITP4915M Group 24 — Furniture ERP System. "
        "The Finance Department module handles payment vouchers (PV), receipt vouchers (RV), "
        "cash-flow dashboards, and outstanding balance reconciliation. "
        "This section describes the implemented WinForms UI aligned with the interim report mock-ups."
    )

    heading(doc, "8.15 Finance Department", level=1)
    numbered(doc, [
        "User can open the Finance Department module from the main navigation menu.",
        "User can switch between four tabs: Dashboard, Payment Vouchers, Receipt Vouchers, and Outstanding.",
        "User can access tabs according to role permissions (Payment Voucher and/or Receipt Voucher view rights).",
        "User can use the module to record supplier payments, customer receipts, review cash flow, and check unsettled PO/invoice balances.",
    ])

    # --- Dashboard ---
    heading(doc, "8.15.1 Dashboard", level=2)
    para(
        doc,
        "The Dashboard tab provides a management overview of confirmed income, recorded expenses, "
        "and net cash flow, all normalised to HKD using document-locked exchange rates at save time."
    )
    numbered(doc, [
        "User can view the Total Income (HKD) summary card — sum of receipt vouchers excluding cancelled status.",
        "User can view the Total Expenses (HKD) summary card — sum of payment vouchers excluding cancelled status.",
        "User can view the Net Cash Flow (HKD) summary card — income minus expenses (green when positive, red when negative).",
        "User can view the Income Trend bar chart — monthly confirmed receipt totals in HKD.",
        "User can view the Income by Payment Method pie chart — receipt breakdown by Cash, Bank Transfer, Credit Card, or Cheque.",
        "User can view the Expense Trend bar chart — monthly payment voucher totals in HKD.",
        "User can view the Expense by Payment Method pie chart — payment breakdown by method.",
        "User can view the Receipt Vouchers by Currency table — foreign total, weighted rate, HKD total, and count per currency.",
        "User can view the Payment Vouchers by Currency table — same breakdown for outgoing payments.",
        "User can click Print Report PDF to export a finance dashboard report including summary fields, charts, and currency breakdown.",
        "User can see dashboard figures refresh automatically when payment or receipt voucher lists are reloaded.",
    ])

    # --- Payment Vouchers list ---
    heading(doc, "8.15.2 Payment Vouchers", level=2)
    para(doc, "The Payment Vouchers tab lists all PV documents and provides create, view, edit, and status-update actions.")
    numbered(doc, [
        "User can click + New Payment Voucher to open the create form (requires create permission).",
        "User can select a row and click View Detail to open a read-only voucher detail dialog.",
        "User can double-click a row to open the edit form when edit permission is granted; otherwise View Detail opens.",
        "User can select a row and click Update Status to change status (Draft, Approved, Paid, Cancelled).",
        "User can use the Payment Voucher Filters block to add multi-condition AND filters on visible list columns.",
        "User can filter by status using the status dictionary (Draft / Approved / Paid / Cancelled).",
        "User can export filtered list results to CSV from the filter block.",
        "User can clear filter conditions or add/remove filter rows.",
        "User can view the payment voucher grid showing voucher code, supplier, amount, currency, HKD equivalent, method, status, and related fields.",
    ])

    heading(doc, "8.15.2.1 Create / Edit Payment Voucher", level=3)
    para(
        doc,
        "The create/edit dialog is split into a scrollable header (voucher fields) and a lower "
        "Purchase Order Allocations section. On save, the sum of Pay Amount lines must equal the header Amount."
    )
    numbered(doc, [
        "User can view Voucher Code — auto-generated on create; editable on edit.",
        "User can search and select Supplier * using type-ahead (minimum 2 characters) or browse the supplier dropdown.",
        "User can view Staff — current logged-in user (read-only label).",
        "User can select Currency * from the currency list.",
        "User can enter Amount * — total payment for this voucher.",
        "User can view HKD Equivalent — system-calculated base amount and exchange rate.",
        "User can select Payment Method — Cash, Bank Transfer, Credit Card, or Cheque.",
        "User can enter Method Ref — optional payment reference (e.g. bank transaction ID).",
        "User can select Status — Draft, Approved, Paid, or Cancelled.",
        "User can enter Remark — optional multi-line notes.",
        "User can view the allocation balance line: Total | Allocated | Remaining (balanced when remaining is zero).",
        "User can click Add Line to add a PO allocation row.",
        "User can click Remove Line to delete the selected allocation row (at least one line required).",
        "User can select Purchase Order * per line from POs belonging to the chosen supplier.",
        "User can enter Pay Amount * per line — positive decimal amount allocated to that PO.",
        "User can select Payment Type * per line — clearing type (e.g. Deposit, Partial, Final) from dictionary.",
        "User can save only when every line has a PO, positive pay amount, and payment type, with no duplicate PO per voucher.",
        "User can save only when allocated total equals the header Amount (within 0.01 tolerance).",
        "User can click Cancel to close without saving.",
        "User can click Save to persist the voucher and allocation lines.",
        "User cannot edit PO allocations when voucher status is Paid or Cancelled (read-only allocation grid).",
    ])

    heading(doc, "8.15.2.2 View Payment Voucher Detail", level=3)
    numbered(doc, [
        "User can view voucher header fields on the Voucher tab (field-value grid).",
        "User can switch to the Supplier tab to view supplier profile and raw-material quote list.",
        "User can switch to the PO Allocations tab to view linked purchase orders, pay amounts, and clearing types.",
        "User can switch to the Related Documents tab when linked documents exist.",
        "User can switch to the Activity tab to view the document audit log.",
        "User can click Close to exit the detail dialog.",
    ])

    heading(doc, "8.15.2.3 Update Payment Voucher Status", level=3)
    numbered(doc, [
        "User can view the current status of the selected payment voucher.",
        "User can select a new status from Draft, Approved, Paid, or Cancelled.",
        "User can click Save to apply the status change.",
        "User can click Cancel to close without changes.",
    ])

    # --- Receipt Vouchers ---
    heading(doc, "8.15.3 Receipt Vouchers", level=2)
    para(doc, "The Receipt Vouchers tab lists RV documents and supports create, verify, view, edit, and status-update workflows.")
    numbered(doc, [
        "User can click + New Receipt Voucher to open the create form (requires create permission).",
        "User can select a row and click View Detail for read-only access.",
        "User can double-click a row to edit when permitted.",
        "User can select a draft receipt and click Verify Receipt to confirm payment and apply invoice allocations.",
        "User can select a row and click Update Status to change status (Draft, Confirmed, Cancelled).",
        "User can use the Receipt Voucher Filters block for multi-condition filtering and CSV export.",
        "User can view the receipt voucher grid with code, customer, amount, currency, method, received date, and status.",
    ])

    heading(doc, "8.15.3.1 Create / Edit Receipt Voucher", level=3)
    para(
        doc,
        "The create/edit dialog has a header section and an Invoice Allocations grid below. "
        "Allocated amounts must sum to the receipt Amount before save."
    )
    numbered(doc, [
        "User can view Voucher Code — auto-generated on create.",
        "User can search and select Customer * using type-ahead customer picker.",
        "User can view Staff — current user (read-only).",
        "User can select Currency * and enter Amount *.",
        "User can view HKD Equivalent — calculated from the selected currency rate.",
        "User can select Payment Method and enter Method Ref.",
        "User can set Received Date * using the date picker.",
        "User can select Status — Draft, Confirmed, or Cancelled (Confirmed vouchers restrict status changes on edit).",
        "User can enter Remark.",
        "User can view allocation balance: Total | Allocated | Remaining.",
        "User can click Add Line to add an invoice allocation row.",
        "User can click Remove Line to delete a row (minimum one line).",
        "User can click Add Exchange Loss to add a line for unallocated foreign-exchange difference (no invoice linked).",
        "User can select Invoice * per line from open invoices for the selected customer.",
        "User can enter Allocated Amount * per line.",
        "User can select Clearing Type * — Deposit, Partial, Final, or Exchange Loss.",
        "User can save when allocated total equals receipt Amount and each non-exchange-loss line has a unique invoice.",
        "User can click Cancel or Save.",
        "User cannot edit invoice allocations when the receipt is already Confirmed (read-only grid).",
    ])

    heading(doc, "8.15.3.2 Verify Receipt", level=3)
    para(
        doc,
        "Verify Receipt is a dedicated workflow for draft receipts. It opens an allocation dialog, "
        "validates that allocations sum to the receipt amount, then confirms the voucher and updates invoice payment status."
    )
    numbered(doc, [
        "User can open Verify Receipt from the Receipt Vouchers toolbar for a selected draft voucher.",
        "User can view receipt code, amount, and customer at the top of the verify dialog.",
        "User can allocate the receipt across one or more invoices using the allocation grid.",
        "User can add, remove, or add exchange-loss lines as in the create form.",
        "User can click Verify Receipt to confirm — system checks allocation balance and saves allocations.",
        "User can click Cancel to abort verification.",
        "User receives a warning if the receipt is already verified.",
    ])

    heading(doc, "8.15.3.3 View Receipt Voucher Detail", level=3)
    numbered(doc, [
        "User can view voucher header on the Voucher tab.",
        "User can switch to the Customer tab — profile, contact persons, and delivery addresses.",
        "User can switch to the Invoice Allocations tab — invoice, amount, and clearing type per line.",
        "User can view Related Documents and Activity tabs when available.",
        "User can click Close.",
    ])

    heading(doc, "8.15.3.4 Update Receipt Voucher Status", level=3)
    numbered(doc, [
        "User can view current status and select Draft, Confirmed, or Cancelled.",
        "User can see a hint that confirmed vouchers cannot revert to Draft.",
        "User can save or cancel the status change.",
    ])

    # --- Outstanding ---
    heading(doc, "8.15.4 Outstanding", level=2)
    para(
        doc,
        "The Outstanding tab shows unsettled accounts payable (purchase orders) and accounts receivable (invoices) "
        "to help finance staff decide how much to pay or record before creating PV/RV documents."
    )
    numbered(doc, [
        "User can click Refresh to reload outstanding balances.",
        "User can view the Accounts Payable — Outstanding POs grid in the upper panel.",
        "User can see AP columns: PO Code, Supplier, Currency, PO Total, Paid, Outstanding, Request Delivery, Status.",
        "User can view the Accounts Receivable — Outstanding Invoices grid in the lower panel.",
        "User can see AR columns: Invoice Code, Customer, Currency, Invoice Total, Received, Outstanding, Invoice Date, Status.",
        "User can use this view to identify which POs still need payment and which invoices still need receipt.",
        "User can cross-check amounts before entering Pay Amount or Allocated Amount on new vouchers.",
    ])

    heading(doc, "Screenshot checklist (for final report)", level=2)
    numbered(doc, [
        "8.15.1 — Finance Department → Dashboard tab (summary cards, charts, currency tables).",
        "8.15.2 — Payment Vouchers list with Filter Block; Create Payment Voucher form with PO Allocations.",
        "8.15.2 — View Payment Voucher Detail (Voucher + PO Allocations tabs).",
        "8.15.3 — Receipt Vouchers list; Create Receipt Voucher with Invoice Allocations balanced.",
        "8.15.3 — Verify Receipt dialog.",
        "8.15.4 — Outstanding tab showing AP and AR grids.",
    ])

    doc.save(OUTPUT)
    print("Wrote", OUTPUT)


if __name__ == "__main__":
    build()
