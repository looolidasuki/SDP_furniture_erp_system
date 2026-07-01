"""Section 8.11 Internal Transfer — UI Design (English), ITP report style: User can..."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

OUTPUT = r"C:\Users\user\source\repos\4915M_claude\FurnitureERP\docs\Section8_11_Internal_Transfer_UI_Design_v2.docx"


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.1)
    s.right_margin = Inches(1.1)

    heading(doc, "8.11 Internal Transfer", level=1)

    numbered(doc, [
        "User can open the Internal Transfer module from the main menu.",
        "User can switch to the Issue RM Request tab to issue raw materials against an RM Request Note.",
        "User can switch to the Other Transfer tab to move stock manually between warehouses.",
        "User can open Raw Material Request Notes from Production → RM Requests, then click Issue Materials "
        "to jump to Issue RM Request with the selected note pre-filled.",
    ])

    heading(doc, "8.11.1 Raw Material Request Notes List (Production Entry)", level=2)
    numbered(doc, [
        "User can add filter conditions (column, operator, value) in the RM Request Filters block.",
        "User can apply filters to search request notes with different conditions.",
        "User can use different compare methods to search with values.",
        "User can input the search values in the filter value field.",
        "User can add more conditions for better search results.",
        "User can clear all search conditions.",
        "User can export the search results in CSV format.",
        "User can remove filter conditions.",
        "User can refresh the list of raw material request notes.",
        "User can view the details of the selected request note (View Detail or double-click a row).",
        "User can click Issue Materials to open Internal Transfer with the selected SCR note pre-selected.",
        "User can view the request list showing Request Code, Production Order Code, Status, Request Date, and Required Date.",
        "User can close the dialog without saving changes.",
    ])

    heading(doc, "8.11.2 Issue RM Request", level=2)

    heading(doc, "Step 1 — Filters", level=3)
    numbered(doc, [
        "User can select or type-ahead a Production Order (e.g. PTO-00000085 — SO-00000101) to narrow the RM Request list.",
        "User can select “(All open production orders)” to show all open request notes.",
        "User can select the Inventory WH * (source inventory warehouse, e.g. CN Warehouse Main).",
        "User can select or search the RM Request * (e.g. SCR-00000085 — PTO-00000085 [Draft]).",
        "User can view the paired Production WH name and ID for the selected inventory warehouse (read-only label).",
        "User can change Production Order, Inventory WH, or RM Request; the issue preview will auto-refresh.",
    ])

    heading(doc, "Step 2 — Issue Preview", level=3)
    numbered(doc, [
        "User can review the issue preview grid that auto-refreshes when the request note or warehouse changes.",
        "User can view Raw Material code for each line (e.g. RM-ADH-0001).",
        "User can view Request Qty — quantity requested on the RM Request Note.",
        "User can view Inventory Available — available stock in the selected inventory warehouse.",
        "User can view Inventory Net — net available stock after reservations.",
        "User can view Production On Hand — current stock in the paired production warehouse.",
        "User can view Min Stock — minimum stock level for the material.",
        "User can view Shortage Qty — shortfall (zero when inventory is sufficient); shortage rows may be highlighted.",
    ])

    heading(doc, "Step 3 — Actions", level=3)
    numbered(doc, [
        "User can read the workflow hint: “If shortage exists: Create PO → receive goods → issue materials.”",
        "User can click Issue Materials to transfer stock from Inventory WH to Production WH for all sufficient lines.",
        "User can click Create PO for Shortages to open a purchase order draft pre-filled with shortage quantities.",
    ])

    heading(doc, "8.11.3 Other Transfer", level=2)
    numbered(doc, [
        "User can select Item Type (Raw Material or Product).",
        "User can select From Warehouse * (source warehouse).",
        "User can select To Warehouse * (destination warehouse; must differ from source).",
        "User can click Load Items to list transferable items and available quantities from the source warehouse.",
        "User can enter Transfer Qty on each line (only this column is editable).",
        "User can click Transfer to confirm and move stock after the system validates available quantities.",
        "User can read the hint that only available stock (physical − reserved) can be transferred.",
    ])

    doc.save(OUTPUT)
    print("Wrote", OUTPUT)


if __name__ == "__main__":
    build()
