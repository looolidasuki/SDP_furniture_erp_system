"""Generate complete Section 8 UI Design document (English) aligned with ITP report structure."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

OUTPUT = r"C:\Users\user\source\repos\4915M_claude\FurnitureERP\docs\Section8_UI_Design_Complete.docx"


def add_title(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def add_section(doc, title, level, intro, items):
    add_title(doc, title, level=level)
    if intro:
        add_para(doc, intro)
    if items:
        add_list(doc, items)


# Each entry: (title, level, intro, items)
SECTIONS = [
    ("8. UI Design", 0, (
        "ITP4915M Group 24 — Furniture ERP System (WinForms Desktop Application). "
        "This section describes the user interface for each module view, following the "
        "interim report structure. Descriptions reflect the implemented application on branch v6."
    ), None),

    ("8.1 Login Screen", 1, None, [
        "Username or Email: The user enters the account username or email address.",
        "Password: The user enters the account password (Enter key submits the form).",
        "Login: The user clicks Login after entering credentials; on success the main application window opens.",
        "Error message: The system displays validation, authentication, or database connection errors below the form.",
    ]),

    ("8.2 Main Dashboard (Overview — My Tasks)", 1, None, [
        "Customers: Displays the total number of customers in the system.",
        "Sales Orders: Displays the total number of sales orders.",
        "Invoices: Displays the total number of invoices.",
        "Products: Displays the total number of products.",
        "Month Income: Displays total receipt income for the current month (HKD, when finance permission is granted).",
        "Month Expense: Displays total payment expense for the current month (HKD, when finance permission is granted).",
        "AR Outstanding: Displays total accounts receivable outstanding (when finance permission is granted).",
        "AP Outstanding: Displays total accounts payable outstanding (when finance permission is granted).",
        "My Tasks: The user opens a tab listing open and overdue tasks assigned to the logged-in user; double-click navigates to the document.",
        "Action Items: The user opens a tab listing proceeding sales orders (uninvoiced balance) and purchase orders (unpaid balance).",
        "Permissions: The user views module permissions by department (super user sees all departments; others see their own department).",
        "Task grid: Displays the list for the selected overview tab.",
    ]),

    ("8.2.1 Main Dashboard (Overview — Action Items)", 2, (
        "Displays proceeding sales orders with uninvoiced balances and purchase orders with unpaid balances. "
        "Double-clicking a row navigates to the corresponding document."
    ), None),

    ("8.2.2 Main Dashboard (Overview — Permissions)", 2, (
        "Displays the permission matrix (View / Create / Edit) for each module under the selected department. "
        "Super users may filter modules by keyword."
    ), None),

    ("8.3 Customer Dashboard", 1, "Accessed from the Sales module — Customers tab.", [
        "New Customer: The user creates a new customer record.",
        "Refresh: The user reloads the customer list.",
        "View Detail: The user opens a read-only detail dialog for the selected customer.",
        "Edit: The user opens the edit dialog for the selected customer.",
        "Filter Block: The user may add multiple conditions (column, operator, value), apply filters, clear conditions, remove rows, and export results to CSV.",
        "Customer grid: Displays customer records matching the filter criteria.",
        "Double-click: Opens view detail, or edit when the user has edit permission.",
    ]),

    ("8.3.1 Create / Edit Customer", 2, "Tabbed dialog: General | Contact Persons | Delivery Addresses.", [
        "Customer Code: Read-only system code (on edit).",
        "Customer Name *: The user enters the customer company name.",
        "Billing Address: The user enters the billing address.",
        "Payment Term: The user selects the payment term from the dictionary.",
        "Contact Persons tab: Editable grid — Contact Person, Title, Phone, Email.",
        "Delivery Addresses tab: Editable grid — Delivery Address, Contact Person, Phone, Email.",
        "Save / Update: Saves the customer and related contact/address rows.",
        "Close: Closes without saving.",
    ]),

    ("8.3.2 View Customer Detail", 2, None, [
        "Customer tab: Displays profile field-value pairs.",
        "Contact Persons tab: Displays all contact persons.",
        "Delivery Addresses tab: Displays all delivery addresses.",
        "Delivery Notes tab: Lists delivery notes linked to this customer.",
        "Receipt Vouchers tab: Lists receipt vouchers linked to this customer.",
    ]),

    ("8.4 Quotation Dashboard", 1, "Accessed from the Sales module — Quotations tab.", [
        "New Quotation: The user creates a new quotation.",
        "Refresh: Reloads the quotation list.",
        "View Detail: Opens read-only quotation detail.",
        "Edit: Opens the edit quotation dialog.",
        "Convert to SO: Converts the selected quotation to a sales order.",
        "View Products / View BOM: Opens the product catalogue viewer for reference.",
        "Filter Block: Multi-condition filter with status dictionary support; server-side paging (100 records per page).",
        "Next / Previous page: The user navigates paged quotation results.",
        "Quotation grid: Displays matching quotation records.",
    ]),

    ("8.4.1 Create / Edit Quotation", 2, None, [
        "Customer *: The user selects the customer.",
        "Staff: Read-only — current logged-in staff.",
        "Currency *: The user selects invoice currency.",
        "FX / Total: System displays exchange rate and calculated total.",
        "Status / Remark: The user sets quotation status and optional remark.",
        "Line grid: Product, Available Stock, Price, Quantity, Discount — user adds product lines.",
        "Save / Update: Saves quotation header and lines.",
        "Close: Closes without saving.",
    ]),

    ("8.5 Sales Order Dashboard", 1, "Accessed from the Sales module — Sales Orders tab.", [
        "New Sales Order: The user creates a new sales order.",
        "Refresh: Reloads the sales order list.",
        "View Detail / Edit: Opens view or edit dialog.",
        "Confirm Order: Confirms the selected sales order.",
        "Create Production Order: Creates a production order from the selected sales order.",
        "Open Orders / Uninvoiced / Show All: Quick list filters.",
        "View Products / View BOM: Product catalogue reference.",
        "Filter Block: Multi-condition filter with sales order status; server-side paging.",
        "Sales order grid: Displays matching records.",
    ]),

    ("8.5.1 Create / Edit Sales Order", 2, None, [
        "Customer *: The user selects the customer.",
        "Staff: Read-only assigned staff.",
        "Currency *: The user selects currency.",
        "Delivery Address *: The user selects or enters delivery address.",
        "Requested Delivery Date: The customer requested delivery date.",
        "Discount: Order-level discount (percentage or fixed).",
        "Customer Ref Number: External PO reference (PO-PL-#########).",
        "Status / Remark: Order status and remark (status locked on create).",
        "Line grid: Product, Available Stock, Price, Quantity, Discount.",
        "Save / Update / Close: Standard save actions.",
    ]),

    ("8.6 Production Dashboard", 1, "Accessed from the Production module.", [
        "New Production Order: Creates a standard production order linked to a sales order.",
        "Quick Entry: Batch-creates production orders from multiple open sales orders.",
        "Sample Production: Creates a sample/trial production order without a sales order.",
        "Refresh: Reloads the production order list.",
        "View Detail / Edit: Opens view or edit dialog.",
        "Complete & Stock In: Completes production and posts finished goods to inventory warehouse.",
        "View Products / Add Product: Product catalogue management.",
        "RM Requests / Create RM Request: Lists or creates raw material request notes for the selected PO.",
        "Filter Block: Production order status filter.",
        "Production grid: Displays production orders.",
    ]),

    ("8.6.1 Create / Edit Production Order", 2, None, [
        "Order Code: Production order code (auto-generated on create).",
        "Sales Order *: Linked sales order (standard PO only).",
        "Staff *: Assigned production staff.",
        "Est. Finish Date *: Estimated completion date.",
        "Status / Remark: Production status and notes.",
        "Product lines grid: Products and quantities to manufacture.",
        "Save / Update / Close.",
    ]),

    ("8.6.2 New Material Request Form", 2, None, [
        "Request Date *: Date of the material request.",
        "Remark: Optional notes for procurement.",
        "Line grid: Raw material, requested quantity — user adds lines for materials needed.",
        "Submit: Creates the raw material request note.",
        "Close: Closes without saving.",
    ]),

    ("8.7 Raw Materials", 1, "Sub-tabs: Materials | Warehouse Stock | Supplier Quotes.", [
        "Materials tab — New Raw Material: Creates a new raw material master record.",
        "Materials tab — Edit: Edits selected raw material (Code, Category, Size, Color, Min Stock Level, Status).",
        "Materials tab — Filter Block: Filters the raw material list.",
        "Warehouse Stock tab — Warehouse selector: Filters stock by warehouse.",
        "Warehouse Stock tab — Grid: Physical, reserved, and available quantities per raw material.",
        "Supplier Quotes tab — Grid: Supplier pricing quotes per raw material line.",
        "Refresh: Reloads the active sub-tab grid.",
    ]),

    ("8.8 Purchase Order Dashboard", 1, "Accessed from the Procurement module — Purchase Orders tab.", [
        "New Purchase Order: Creates a new PO.",
        "Refresh: Reloads the PO list.",
        "View Detail / Edit: Opens view or edit dialog.",
        "Unpaid POs / Show All: Quick list filters.",
        "Filter Block: Multi-condition filter with PO status; server-side paging.",
        "Purchase order grid: Displays matching PO records.",
    ]),

    ("8.8.1 Create Purchase Order", 2, None, [
        "Supplier *: The user selects the supplier.",
        "Billing Address: Supplier billing address (from supplier master).",
        "Receiving Warehouse *: Warehouse where goods will be received.",
        "Ship-To Address *: Delivery destination address.",
        "Bill-To Address: Billing entity address.",
        "Buyer (Staff): Purchasing staff responsible.",
        "Currency *: PO currency.",
        "Request Delivery Date *: Expected delivery date.",
        "Status / Remark: PO status and notes.",
        "Line grid: Raw material lines — material, price, order quantity.",
        "Save / Close.",
    ]),

    ("8.8.2 Edit / View Purchase Order", 2, None, [
        "PO Code: Auto-generated purchase order code.",
        "All header fields as create form; supplier lock may apply after receipt activity.",
        "Financial Summary: Total amount display on edit.",
        "View Detail tabs: Order Lines | Payment Vouchers | Goods Received | Print PDF.",
    ]),

    ("8.9 Goods Received Dashboard", 1, "Accessed from the Procurement module — Goods Received tab.", [
        "New GRN: Creates a new Goods Received Note.",
        "Refresh: Reloads the GRN list.",
        "View Detail: Opens read-only GRN detail.",
        "Edit: Opens the edit GRN dialog.",
        "Confirm Receipt: Confirms receipt and posts stock to warehouse.",
        "Filter Block: Multi-condition filter on GRN list.",
        "Export CSV: Exports filtered results.",
        "GRN grid: Displays matching goods received notes.",
    ]),

    ("8.9.1 Create Goods Receipt Note", 2, None, [
        "Purchase Order *: The user selects the related purchase order.",
        "Supplier: System displays the supplier name from the selected PO.",
        "Remark: Optional user remark.",
        "Line grid: Raw materials on the PO — user enters actual received quantity per line.",
        "Save / Close.",
    ]),

    ("8.9.2 View Goods Received Note", 2, None, [
        "Detail tab: GRN header fields and received line quantities.",
        "Activity tab: Document audit log of changes.",
        "Print PDF: Exports the GRN to PDF.",
    ]),

    ("8.9.3 Goods Received Note Activity", 2, (
        "Activity tab within View GRN — displays timestamped change log entries for the document."
    ), None),

    ("8.9.4 Edit Goods Received Note", 2, None, [
        "GRN Code: Auto-generated code (read-only).",
        "Purchase Order *: User may change PO if no receipt posted.",
        "Supplier: Display supplier of selected PO.",
        "Status: User may change GRN status.",
        "Remark: User may update remark.",
        "Line grid: User may update actual received quantities.",
        "Close: Closes without save.",
        "Save: Updates GRN data.",
    ]),

    ("8.10 Warehouse Dashboard", 1, "Tabbed: Warehouses | Delivery Notes | Inventory Ledger.", [
        "Warehouses tab — New Warehouse: Creates a warehouse.",
        "Warehouses tab — Refresh: Reloads warehouse list.",
        "Warehouses tab — View Stock: Opens stock detail for selected warehouse.",
        "Warehouses tab — Edit: Edits warehouse name and address.",
        "Warehouses tab — Filter Block: Multi-condition warehouse filter.",
        "Stock preview pane: Selecting a warehouse shows Products and Raw Materials sub-tabs with on-hand stock.",
        "Delivery Notes tab — shortcut: Same as Section 8.12.",
        "Inventory Ledger tab — shortcut: Same as Section 8.12.4.",
    ]),

    ("8.10.1 Create New Warehouse", 2, None, [
        "Warehouse Type *: User selects warehouse type (Inventory / Production).",
        "Warehouse Name *: User enters warehouse name.",
        "Address: User enters warehouse address.",
        "Paired production warehouse: System may auto-create a paired warehouse when type is Inventory.",
        "Cancel: Closes without save.",
        "Save: Creates the new warehouse.",
    ]),

    ("8.10.2 View Stock from Warehouse", 2, None, [
        "Selected warehouse: Displays the warehouse being viewed.",
        "Products tab: Lists product stock (physical, reserved, available) in the warehouse.",
        "Raw Materials tab: Lists raw material stock in the warehouse.",
        "Grid: Displays item code, quantities, and alert flags for low stock.",
    ]),

    ("8.11 Internal Transfer", 1, (
        "Accessed from the Warehouse / Internal Transfer module. Two workflows: "
        "request-based material issue and free warehouse-to-warehouse transfer."
    ), None),

    ("8.11.1 Issue RM Request", 2, None, [
        "Production Order: User selects or searches a production order to narrow RM request notes.",
        "Inventory WH *: User selects the inventory warehouse that supplies materials.",
        "RM Request *: User selects or searches an open Raw Material Request Note.",
        "Production WH: System displays the paired production warehouse (read-only).",
        "Issue preview grid: Auto-refreshes showing Request Qty, Inventory Available, Production On-hand, and Shortage per line.",
        "Issue Materials: Posts stock from inventory warehouse to production warehouse.",
        "Create PO for Shortages: Opens a purchase order draft pre-filled with shortage quantities.",
    ]),

    ("8.11.2 Other Transfer (Free Transfer)", 2, None, [
        "Item Type: User selects Raw Material or Product.",
        "From Warehouse * / To Warehouse *: Source and destination (must differ).",
        "Load Items: Loads transferable items with available quantities from source warehouse.",
        "Transfer Qty: User enters quantity per line (only this column is editable).",
        "Transfer: Validates and executes the stock move; logs Transfer In/Out to Inventory Ledger.",
    ]),

    ("8.12 Delivery Notes Dashboard", 1, "Warehouse module — Delivery Notes tab.", [
        "New Delivery Note: Creates a delivery note.",
        "Refresh: Reloads the delivery note list.",
        "View Detail: Opens read-only detail with tabs.",
        "Confirm Delivery: Confirms shipment and posts stock deduction.",
        "Edit: Opens edit dialog.",
        "Print DN PDF / Print Reply Slip: Exports delivery note or customer reply slip.",
        "Filter Block: Multi-condition filter with delivery status dictionary.",
        "Delivery note grid: Displays matching records.",
    ]),

    ("8.12.1 Create New Delivery Note", 2, None, [
        "Customer *: User selects customer.",
        "Sales Order *: User selects linked sales order.",
        "Warehouse *: Shipment warehouse.",
        "Ship Method *: Delivery method from dictionary.",
        "Tracking Number: Carrier tracking reference.",
        "Status: Delivery note status.",
        "Staff: Approving / dispatch staff.",
        "Signed By / Signed Date: Customer sign-off (may be filled on confirm).",
        "Remark: Optional notes.",
        "Product line grid: User edits actual ship quantity per product line.",
        "Cancel / Save.",
    ]),

    ("8.12.2 View Delivery Note", 2, None, [
        "Detail tab: Delivery note header and product lines.",
        "Activity tab: Document change log.",
        "Print DN PDF / Print Reply Slip.",
    ]),

    ("8.12.3 Edit Delivery Note", 2, None, [
        "Delivery Note Code / Reply Slip Code: Auto-generated identifiers.",
        "Customer, Sales Order, Warehouse, Ship Method, Tracking Number: Editable header fields.",
        "Status, Signed By, Signed Date, Remark: Editable.",
        "Product line grid: User may change ship quantities.",
        "Cancel / Save.",
    ]),

    ("8.12.4 Inventory Ledger", 1, (
        "Read-only audit trail of all stock movements. Accessed from Warehouse → Inventory Ledger tab."
    ), [
        "Refresh: Reloads the latest ledger entries (up to 300 rows).",
        "Warehouse filter: All Warehouses or a specific warehouse.",
        "Item Type filter: All Types, Raw Material, or Product.",
        "Ledger grid columns: Date, Item Type, Item, Warehouse, Qty Change, Balance After, Action, Document Type, Document Code, Staff, Remark.",
        "Automatic logging: Entries are created when GRN, Delivery Note, Internal Transfer, or material issue operations complete.",
        "Read-only: Quantities cannot be edited on this screen.",
    ]),

    ("8.13 Invoice Dashboard", 1, "Finance module — Invoices tab.", [
        "Invoices / Refunds: User switches between invoice and refund request views.",
        "New Invoice: Creates a new invoice.",
        "Refresh: Reloads the list.",
        "View Detail / Edit: Opens invoice view or edit.",
        "Print PDF: Exports selected invoice.",
        "Invoice from Delivery: Quick-create invoice from an uncompleted delivery note.",
        "Filter Block: Multi-condition filter; server-side paging (100 records per page).",
        "Next / Previous page: Navigates paged results.",
        "Invoice grid: Displays matching invoices.",
    ]),

    ("8.13.1 Create New Invoice", 2, None, [
        "Customer *: User selects customer.",
        "Sales Order *: User selects related sales order.",
        "Total Amount: System displays sales order amount.",
        "Invoice Type *: Deposit or standard invoice.",
        "Amount / Line grid: Deposit amount or product lines depending on type.",
        "Staff: Creating staff (read-only).",
        "Status / Remark: Invoice status and notes.",
        "Cancel / Save.",
    ]),

    ("8.13.2 View Invoice", 2, None, [
        "Invoice details: Header field-value display.",
        "Related Documents tab: Linked sales order and delivery notes.",
        "Activity tab: Document audit log.",
        "Deposit lines / Invoice lines: Line detail grid.",
        "Print PDF.",
    ]),

    ("8.13.2.1 Related Document", 3, None, [
        "Displays the sales order and delivery notes related to this invoice.",
        "User may select a related document to print.",
    ]),

    ("8.13.2.2 Activity", 3, (
        "Displays timestamped change log entries for the invoice document."
    ), None),

    ("8.13.3 Edit Invoice", 2, None, [
        "Invoice Code: Auto-generated reference.",
        "Customer / Sales Order: Editable when permitted.",
        "Staff: Creating staff.",
        "Invoice Type / Deposit Amount: Editable for deposit invoices.",
        "Status / Remark: Editable.",
        "Financial Summary: System-calculated totals.",
        "Line grid: Invoice line details.",
        "Close / Update.",
    ]),

    ("8.13.4 Invoice From Delivery", 2, None, [
        "Delivery note picker: User selects an uncompleted delivery note.",
        "Cancel: Closes without save.",
        "Create: Creates a new invoice from the selected delivery note.",
    ]),

    ("8.14 Refund Request Dashboard", 1, "Finance module — Refunds tab.", [
        "New Refund: Creates a refund request note.",
        "Refresh: Reloads the list.",
        "View Detail: Opens multi-tab refund detail dialog.",
        "Update Status: Changes refund request status.",
        "Edit: Opens edit dialog.",
        "Print PDF: Exports refund request to PDF.",
        "Filter Block: Multi-condition filter with refund status dictionary.",
        "Refund grid: Displays matching refund requests.",
    ]),

    ("8.14.1 Create New Refund Request Note", 2, None, [
        "Related Invoice *: User enters or selects invoice code.",
        "Invoice Total: System displays invoice total amount.",
        "Related Receipt Voucher: User links the receipt voucher being refunded.",
        "Staff: Creating staff (read-only).",
        "Refund Amount *: User enters refund amount.",
        "Refund Method * / Refund Reason *: Dictionary selections.",
        "Refund Transaction Ref: Bank/payment reference.",
        "Status / Remark: Refund status and notes.",
        "Cancel / Save.",
    ]),

    ("8.14.2.1 View Refund Request — Request Details", 2, None, [
        "Request tab: Refund request header field-value pairs.",
        "Customer tab: Customer profile with contact persons and delivery addresses.",
        "Invoice tab: Related invoice financial context (see 8.14.2.3).",
        "Receipt Voucher tab: Linked receipt voucher detail and allocations.",
        "Activity tab: Document audit log.",
        "Close / Print PDF.",
    ]),

    ("8.14.2.2 View Refund Request — Customer Details", 2, None, [
        "Customer information: Company profile of the refund customer.",
        "Contact Persons tab: Contact person details.",
        "Delivery Addresses tab: Delivery address list.",
        "Close / Print PDF.",
    ]),

    ("8.14.2.3 View Refund Request — Related Invoice", 2, (
        "Shows the financial context of the invoice linked to the refund request."
    ), [
        "Invoice header (upper section): Invoice Code, Customer, Sales Order, Invoice Type, Status, Amount, Currency, Remark.",
        "Deposit Line / Invoice Lines tab: Tab title switches by invoice type — Deposit Line for deposit invoices; Invoice Lines for standard invoices.",
        "Receipts & Refunds tab: Lists receipt vouchers and settlement rows tied to this invoice, including allocation type (Deposit, Partial, Final).",
        "Close: Closes the dialog (read-only view).",
        "Print PDF: Exports refund request detail.",
    ]),

    ("8.14.2.4 View Refund Request — Activity", 2, None, [
        "Displays the changes logged for the refund request document.",
        "Close / Print PDF.",
    ]),

    ("8.14.3 Edit Refund Request Note", 2, None, [
        "Refund Request Code: Auto-generated.",
        "Related Invoice / Invoice Amount: Editable invoice link.",
        "Staff: Creating staff.",
        "Related Receipt Voucher: Editable.",
        "Refund Amount, Method, Reason, Transaction Ref: Editable.",
        "Status / Remark: Editable.",
        "Cancel / Save.",
    ]),

    ("8.15 Finance Department", 1, (
        "Dedicated finance module for payment vouchers, receipt vouchers, dashboards, and reconciliation."
    ), None),

    ("8.15.1 Finance Dashboard", 2, None, [
        "Total Income (HKD): Sum of confirmed receipt vouchers.",
        "Total Expenses (HKD): Sum of payment vouchers.",
        "Net Cash Flow (HKD): Income minus expenses.",
        "Income / Expense trend charts: Bar charts by month.",
        "Payment method pie charts: Breakdown by method.",
        "Receipt Vouchers by Currency / Payment Vouchers by Currency: Currency breakdown tables.",
        "Print Report PDF: Exports dashboard summary.",
    ]),

    ("8.15.2 Payment Vouchers", 2, None, [
        "User can click + New Payment Voucher to create a PV with PO allocation lines.",
        "User can click View Detail for read-only voucher, supplier profile, and PO allocations.",
        "User can double-click a row to edit when permitted.",
        "User can click Update Status to change Draft / Approved / Paid / Cancelled.",
        "User can filter the PV list with the Filter Block (multi-condition AND, CSV export).",
        "Create/Edit header: Voucher Code, Supplier * (type-ahead), Staff, Currency *, Amount *, HKD Equivalent, Payment Method, Method Ref, Status, Remark.",
        "PO Allocations: Purchase Order *, Pay Amount *, Payment Type *; Add Line / Remove Line.",
        "User can view Total | Allocated | Remaining balance; save requires allocated total = Amount.",
        "Paid/Cancelled vouchers: allocation grid is read-only.",
    ]),

    ("8.15.3 Receipt Vouchers", 2, None, [
        "User can click + New Receipt Voucher to create an RV with invoice allocation lines.",
        "User can click Verify Receipt on a draft to confirm and apply allocations.",
        "User can click View Detail / Update Status; double-click to edit when permitted.",
        "User can filter the RV list with the Filter Block.",
        "Create/Edit header: Voucher Code, Customer * (type-ahead), Staff, Currency *, Amount *, HKD Equivalent, Payment Method, Method Ref, Received Date *, Status, Remark.",
        "Invoice Allocations: Invoice *, Allocated Amount *, Clearing Type *; Add Line, Remove Line, Add Exchange Loss.",
        "User can view allocation balance; save requires allocated total = receipt Amount.",
        "Confirmed receipts: allocation grid is read-only on edit.",
    ]),

    ("8.15.4 Outstanding (Reconciliation)", 2, None, [
        "User can click Refresh to reload outstanding data.",
        "Accounts Payable — Outstanding POs: PO Code, Supplier, Currency, PO Total, Paid, Outstanding, Request Delivery, Status.",
        "Accounts Receivable — Outstanding Invoices: Invoice Code, Customer, Currency, Invoice Total, Received, Outstanding, Invoice Date, Status.",
        "User can use this tab to identify unsettled POs and invoices before creating PV or RV.",
    ]),

    ("8.16 Supplier Dashboard", 1, "Procurement module — Suppliers tab.", [
        "New Supplier: Creates a supplier record.",
        "Refresh: Reloads supplier list.",
        "View Detail: Opens supplier profile and raw material supply list.",
        "Edit: Opens edit dialog.",
        "Filter Block: Multi-condition supplier filter with CSV export.",
        "Supplier grid: Displays matching suppliers.",
    ]),

    ("8.16.1 Create New Supplier", 2, None, [
        "Supplier Name *, Contact Person, Phone, Email, Billing Address, Payment Term, Status.",
        "Cancel / Save.",
    ]),

    ("8.16.2 View Supplier Details", 2, None, [
        "Supplier profile: Field-value header display.",
        "Raw materials supplied: Grid of materials linked to this supplier.",
    ]),

    ("8.16.3 Edit Supplier", 2, None, [
        "All create fields editable.",
        "Cancel / Save.",
    ]),

    ("8.17 Staff Dashboard", 1, "Staff Administration module (super user).", [
        "New Staff: Creates a staff account.",
        "Refresh: Reloads staff list.",
        "Reset Password: Opens password reset dialog for selected staff.",
        "Edit: Opens edit staff dialog.",
        "Filter Block: Multi-condition staff filter.",
        "Staff grid: Displays staff records.",
    ]),

    ("8.17.1 Create New Staff", 2, None, [
        "Username *, First Name *, Last Name *, Title *, Department *, Email *, Password *, Phone.",
        "Close / Save.",
    ]),

    ("8.17.2 Reset Password", 2, None, [
        "New Password *: User enters new password.",
        "Confirm Password *: User re-enters for confirmation.",
        "Cancel / Save.",
    ]),

    ("8.17.3 Edit Staff", 2, None, [
        "Username, First Name, Last Name, Title, Department, Email, Status, Phone — editable.",
        "Close / Save.",
    ]),

    ("8.18 System Administration", 1, "System Admin module (super user).", [
        "Import CSV: Opens master data import dialog (Customer, Supplier, Product, Raw Material).",
        "System Dictionary: Read-only dictionary grid with filter block.",
        "Product Catalog: Product list with stock alert columns and filter block.",
        "Currency & Exchange Rates: Currency grid with Edit Rate and Refresh.",
        "Filter Block (per card): Multi-condition filter, CSV export on applicable grids.",
    ]),

    ("8.18.1 Import CSV", 2, None, [
        "Import target: User selects table type (Customer, Supplier, Product, Raw Material).",
        "Browse: User selects CSV file to import.",
        "Sample folder: User opens sample data folder for reference CSV templates.",
        "Load sample: Previews sample file for selected target.",
        "Upsert checkbox: Update existing records when key matches.",
        "Import: Executes CSV import.",
    ]),

    ("8.18.2 Currency & Exchange Rates", 2, None, [
        "Currency grid: Currency code, rate to HKD base, enabled flag.",
        "Edit Rate: User updates exchange rate and note (documents lock rate at save time).",
        "Refresh: Reloads currency list.",
        "HKD is the system base currency.",
    ]),
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.1)
    sec.right_margin = Inches(1.1)

    for entry in SECTIONS:
        title, level, intro, items = entry
        add_section(doc, title, level, intro, items)

    doc.save(OUTPUT)
    print("Wrote", OUTPUT)


if __name__ == "__main__":
    build()
