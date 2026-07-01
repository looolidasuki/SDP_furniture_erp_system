"""Generate Section 8 UI Design supplement (English) as DOCX."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

OUTPUT = r"C:\Users\user\source\repos\4915M_claude\FurnitureERP\docs\Section8_UI_Design_Supplement.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("Section 8 — UI Design (Supplement)", level=0)
    for run in title.runs:
        run.font.name = "Calibri"

    add_body(
        doc,
        "ITP4915M Group 24 — Furniture ERP System. "
        "This document completes the UI Design subsections marked for clarification in the interim report, "
        "aligned with the implemented WinForms application."
    )

    # 8.11
    add_heading(doc, "8.11 Internal Transfer", level=1)
    add_body(
        doc,
        "The Internal Transfer module supports stock movement between warehouses and issuing raw materials "
        "to production based on Raw Material Request Notes. The screen is organised as a tabbed view with "
        "two workflows."
    )

    add_heading(doc, "8.11.1 Issue RM Request (Request-based issue)", level=2)
    add_numbered_list(doc, [
        "Production Order: The user may select or type-ahead a production order to narrow the list of open RM request notes.",
        "Inventory WH *: The user selects the inventory warehouse that will supply the materials.",
        "RM Request *: The user selects or searches for a Raw Material Request Note; the picker is filtered by the selected production order when applicable.",
        "Production WH: The system displays the paired production warehouse for the selected inventory warehouse (read-only label).",
        "Issue preview grid: When the request note or warehouse changes, the system automatically refreshes a read-only preview showing Request Qty, Inventory Available, Production On-hand, and Shortage per line.",
        "Issue Materials: On confirmation, the system deducts stock from the inventory warehouse and adds it to the production warehouse; quantities cannot exceed available stock.",
        "Create PO for Shortages: If any line shows a shortage, the user may open a shortcut dialog to create a purchase order draft pre-filled with shortage quantities.",
        "Successful issues are recorded in the Inventory Ledger (see Section 8.12.4).",
    ])

    add_heading(doc, "8.11.2 Other Transfer (Free transfer)", level=2)
    add_numbered_list(doc, [
        "Item Type: The user selects whether to transfer Raw Material or Product stock.",
        "From Warehouse *: The user selects the source warehouse.",
        "To Warehouse *: The user selects the destination warehouse; source and destination must be different.",
        "Load Items: The system loads transferable items from the source warehouse, including Item Code and Available Qty.",
        "Transfer Qty: The user enters the quantity to move on each line; only the Transfer Qty column is editable.",
        "Transfer: On confirmation, the system validates that each transfer quantity does not exceed available stock and posts the movement.",
        "After transfer, the item list is refreshed; physical stock changes are logged as Transfer Out and Transfer In in the Inventory Ledger.",
        "This tab performs an immediate stock move and does not create a separate transfer document number (distinct from request-based issuing).",
    ])

    # 8.12.4
    add_heading(doc, "8.12.4 Inventory Ledger", level=1)
    add_body(
        doc,
        "The Inventory Ledger is a read-only audit view of all stock quantity changes. "
        "It is accessed from the Warehouse module via the Inventory Ledger tab (alongside Warehouse and Delivery Notes)."
    )
    add_numbered_list(doc, [
        "Refresh: The user clicks Refresh to reload the latest stock movement entries.",
        "Warehouse: The user may filter by a specific warehouse or choose All Warehouses.",
        "Item Type: The user may filter by All Types, Raw Material, or Product.",
        "Ledger grid: Displays up to the 300 most recent entries, including Date, Item Type, Item (code), Warehouse, Qty Change (positive = in, negative = out), Balance After, Action, Document Type, Document Code, Staff, and Remark.",
        "Entries are created automatically when operations such as Goods Received, Delivery Note shipment, Internal Transfer, or material issue are completed; users do not enter ledger rows manually.",
        "This view is used for stock reconciliation and traceability; quantities cannot be edited on this screen.",
    ])

    # 8.14.2.3
    add_heading(doc, "8.14.2.3 View Refund Request Note — Related Invoice", level=1)
    add_body(
        doc,
        "When viewing a Refund Request, the Related Invoice tab shows the full financial context of the linked invoice "
        "so finance staff can verify which invoice is being refunded, how much has been received, and whether the invoice is a deposit invoice."
    )
    add_body(doc, "The layout is split into an upper header area and a lower tabbed detail area.")
    add_numbered_list(doc, [
        "Invoice header (upper section): Displays linked invoice header fields such as Invoice Code, Customer, Sales Order, Invoice Type, Status, Amount, Currency, and Remark.",
        "Deposit Line / Invoice Lines tab (lower section): Tab title depends on invoice type — Deposit Line for deposit invoices; Invoice Lines for standard invoices showing line-level product or amount detail.",
        "Receipts & Refunds tab: Lists receipt vouchers and related settlement rows tied to this invoice, including voucher code, date, settled amount, and allocation type (e.g. Deposit, Partial, Final).",
        "Close: Closes the dialog without saving (read-only view).",
        "Print PDF: Exports the refund request detail view to PDF.",
    ])
    add_body(
        doc,
        "Business note: Finance users rely on this tab to confirm that the refund request references the correct invoice, "
        "that prior receipts support the refund amount, and whether the invoice represents a deposit or a goods invoice."
    )

    # 8.15
    add_heading(doc, "8.15 Finance Department", level=1)
    add_body(
        doc,
        "The Finance Department module handles payment vouchers (PV), receipt vouchers (RV), "
        "cash-flow dashboards, and outstanding balance reconciliation. "
        "The interface uses a tabbed layout: Dashboard, Payment Vouchers, Receipt Vouchers, and Outstanding."
    )

    add_heading(doc, "8.15.1 Dashboard", level=2)
    add_numbered_list(doc, [
        "The user can view Total Income (HKD), Total Expenses (HKD), and Net Cash Flow (HKD) summary cards.",
        "The user can view income and expense trend bar charts (monthly HKD totals).",
        "The user can view income and expense pie charts broken down by payment method.",
        "The user can view Receipt Vouchers by Currency and Payment Vouchers by Currency breakdown tables.",
        "The user can click Print Report PDF to export the dashboard summary with charts.",
    ])

    add_heading(doc, "8.15.2 Payment Vouchers", level=2)
    add_numbered_list(doc, [
        "The user can create a new payment voucher (+ New Payment Voucher).",
        "The user can view the list, double-click to edit (with permission), or use View Detail for read-only access.",
        "The user can Update Status (Draft, Approved, Paid, Cancelled).",
        "The user can use the Filter Block for multi-condition filtering and CSV export.",
        "Create/Edit form — header: Voucher Code, Supplier *, Staff, Currency *, Amount *, HKD Equivalent, Payment Method, Method Ref, Status, Remark.",
        "Purchase Order Allocations grid: Purchase Order *, Pay Amount *, Payment Type *; Add Line / Remove Line.",
        "The sum of Pay Amount lines must equal the voucher Amount; balance line shows Total | Allocated | Remaining.",
        "View Detail — tabs: Voucher, Supplier profile, PO Allocations, Related Documents, Activity.",
    ])

    add_heading(doc, "8.15.3 Receipt Vouchers", level=2)
    add_numbered_list(doc, [
        "The user can create a new receipt voucher (+ New Receipt Voucher).",
        "The user can Verify Receipt to confirm a draft and apply invoice allocations.",
        "The user can View Detail, Update Status (Draft, Confirmed, Cancelled), and edit via double-click.",
        "The user can use the Filter Block on the receipt voucher list.",
        "Create/Edit form — header: Voucher Code, Customer *, Staff, Currency *, Amount *, HKD Equivalent, Payment Method, Method Ref, Received Date *, Status, Remark.",
        "Invoice Allocations grid: Invoice *, Allocated Amount *, Clearing Type *; Add Line, Remove Line, Add Exchange Loss.",
        "Allocated total must equal receipt Amount; exchange-loss lines have no invoice.",
        "View Detail — tabs: Voucher, Customer profile, Invoice Allocations, Related Documents, Activity.",
    ])

    add_heading(doc, "8.15.4 Outstanding", level=2)
    add_numbered_list(doc, [
        "The user can refresh unsettled accounts payable and receivable balances.",
        "Accounts Payable grid: outstanding purchase orders with PO Total, Paid, and Outstanding per PO.",
        "Accounts Receivable grid: outstanding invoices with Invoice Total, Received, and Outstanding per invoice.",
        "Used to validate balances before creating payment or receipt vouchers.",
    ])

    add_heading(doc, "Screenshot checklist (for final report)", level=2)
    add_numbered_list(doc, [
        "8.11 — Internal Transfer: both Issue RM Request and Other Transfer tabs.",
        "8.12.4 — Warehouse module → Inventory Ledger tab.",
        "8.14.2.3 — Refund Request View → Invoice tab with Deposit/Invoice Lines and Receipts & Refunds.",
        "8.15 — Finance Department: Dashboard, Payment Vouchers, Receipt Vouchers, and Outstanding tabs.",
    ])

    doc.save(OUTPUT)
    print("Wrote", OUTPUT)


if __name__ == "__main__":
    build()
